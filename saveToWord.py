import os
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

import re
from typing import List, Dict, Any

_FONT_SIZE_NO2 = 22
_FONT_SIZE_NO3 = 16
_COLOR_RED   = "rgb(255,0,0)"
_COLOR_BLUE  = "rgb(0,0,255)"

# 预编译正则
_RE_SUB   = re.compile(r'^[一二三四五六七八九十]+、')              # 一、
_RE_MINOR = re.compile(r'^[（(][一二三四五六七八九十]+[)）]')      # （一） (二)
_RE_SPLIT = re.compile(
    r'^[（(][一二三四五六七八九十]+[)）]'
    r'([^\s。，！？；：]*[\s。，！？；：]|[^\s。，！？；：]*)'  # 标题段
    r'(.*)$'                                                 # 剩余正文
)  # 拆分小标题


def _patch_style(style: str,
                 font: str | None = None,
                 size: int | None = None,
                 color: str | None = None) -> str:
    if font:
        style = re.sub(r'font-family:[^;]+;?', '', style)
        style = f"{style};font-family:{font}".strip(';')
    if size:
        style = re.sub(r'font-size:\d+px;?', '', style)
        style = f"{style};font-size:{size}px".strip(';')
    if color:
        style = re.sub(r'rgb\(\d+,\s*\d+,\s*\d+\);?', '', style)
        style = f"{style};{color}".strip(';')
    return style


def _split_minor(item: Dict[str, Any]) -> List[Dict[str, Any]]:
    """把（一）标题段[。]正文段 拆成两段，无标点也适用"""
    text  = item['text'].strip()
    style = item['attrs'].get('style', '')

    m = _RE_SPLIT.match(text)
    if not m:
        return [item]

    title_part = text[:m.start(1)] + m.group(1).rstrip()
    body_part  = m.group(2).lstrip()

    title_style = _patch_style(style, font='黑体', size=_FONT_SIZE_NO3, color=_COLOR_BLUE)
    body_style  = _patch_style(style, font='仿宋_GB2312', size=_FONT_SIZE_NO3)

    out = [{'text': title_part, 'attrs': {'style': title_style}}]
    if body_part:                       # 避免正文为空时再插空段落
        out.append({'text': body_part, 'attrs': {'style': body_style}})
    return out


def formated_text_object(text_objects: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """
    1. 第1个元素：方正小标宋简体 二号
    2. 副标题（一、）：黑体 三号 红色
    3. 小标题（（一）/ (二) ）：黑体 三号 蓝色，并自动拆分
    4. 其余：仿宋 三号
    返回新的列表
    """
    new_list: List[Dict[str, Any]] = []

    for idx, item in enumerate(text_objects):
        text  = item.get('text', '').strip()
        if not text:
            continue
        style = item.get('attrs', {}).get('style', '')

        # 规则1：第一个元素
        if idx == 0:
            style = _patch_style(style, font='方正小标宋简体', size=_FONT_SIZE_NO2)
            new_list.append({'text': text, 'attrs': {'style': style}})
            continue

        # 规则2：副标题
        if _RE_SUB.match(text):
            style = _patch_style(style, font='黑体', size=_FONT_SIZE_NO3, color=_COLOR_RED)
            new_list.append({'text': text, 'attrs': {'style': style}})
            continue

        # 规则3：小标题（带正文）→ 拆分
        if _RE_MINOR.match(text):
            new_list.extend(_split_minor(item))
            continue

        # 规则4：正文
        style = _patch_style(style, font='仿宋_GB2312', size=_FONT_SIZE_NO3)
        new_list.append({'text': text, 'attrs': {'style': style}})

    return new_list

def sanitize_filename(path: str) -> str:
    """只替换文件名中的非法字符，保留目录分隔符"""
    # 1. 把路径先拆成目录和文件名
    head, tail = os.path.split(path)
    # 2. 替换文件名中的非法字符
    tail = re.sub(r'[<>:\"|?*\x00-\x1f]', '_', tail)
    # 3. 重新拼回去
    return os.path.join(head, tail) if head else tail

def save_to_word(text_objects, filepath="output.docx"):
    """将带样式的文本保存到Word文档"""
    filepath = sanitize_filename(filepath)
    
    try:
        if os.path.exists(filepath):
            print(f"文件已存在，且不允许覆盖。")
            return False
            
        os.makedirs(os.path.dirname(filepath), exist_ok=True)
        doc = Document()
        
        # 设置默认节格式（如页边距）
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        
        for item in text_objects:
            text = item['text'].strip()
            if not text:
                continue
                
            style = item['attrs'].get('style', '')
            # print("style",style)
            p = doc.add_paragraph(style='Normal')

            # --- 对齐方式 ---
            if 'text-align: center' in style:
                # print("居中")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.line_spacing = Pt(43)  # 行距需要单独设置
            elif 'text-align:right' in style:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
            # --- 字体样式 ---
            run = p.add_run(text)
            font = run.font
            
            font_family_dict = {
                "方正小标宋简体" : "方正小标宋简体常规",
                "仿宋_GB2312":"仿宋_GB2312常规",
                "黑体":"思源黑体"
            }

            # 1. 直接提取 font-family 声明
            fam_match = re.search(r'font-family\s*:\s*([^;]+)', style, re.I)
            if fam_match:
                # 去掉可能的引号 / 空格
                fontName = fam_match.group(1).strip().strip('"\'')
                font.name = font_family_dict[fontName]
                # print("font.name",font.name)
            else:
                font.name = '宋体'
                
            # --- 字号 ---
            size_match = re.search(r'font-size:(\d+)px', style)
            if size_match:
                font.size = Pt(int(size_match.group(1)))
            else:
                font.size = Pt(12)  # 默认字号
                
            # --- 颜色 ---
            color_match = re.search(r'rgb\(\s*(\d+)\s*,\s*(\d+)\s*,\s*(\d+)\s*\)', style)
            if color_match:
                r, g, b = map(int, color_match.groups())
                font.color.rgb = RGBColor(r, g, b)
                
            # --- 段落格式 ---
            para_format = p.paragraph_format
            
            # 行距（精确到磅）
            line_match = re.search(r'line-height\s*:\s*(\d+)\s*px', style, re.IGNORECASE)
            if line_match:
                para_format.line_spacing = Pt(float(line_match.group(1)))
                para_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                
            # 缩进（首行缩进和左右缩进）
            indent_match = re.search(r'text-indent\s*:\s*(\d+)\s*px', style, re.IGNORECASE)
            if indent_match:
                para_format.first_line_indent = Pt(float(indent_match.group(1)))
                
            # 段前段后距
            margin_top = re.search(r'margin-top\s*:\s*(\d+)\s*px', style, re.IGNORECASE)
            margin_bottom = re.search(r'margin-bottom\s*:\s*(\d+)\s*px', style, re.IGNORECASE)
            if margin_top:
                para_format.space_before = Pt(float(margin_top.group(1)))
            if margin_bottom:
                para_format.space_after = Pt(float(margin_bottom.group(1)))
                
            # 文字环绕（简化处理）
            if 'text-wrap-mode: wrap' in style:
                para_format.widow_control = True

        doc.save(filepath)
        print(f"Word文档已生成: {filepath}")
        return True
        
    except Exception as e:
        print(f"文件保存失败: {str(e)}")
        return False

# 使用示例

if __name__ == "__main__":    
    text_objects = [{'text': 'XX市XX航空分公司深入贯彻中央八项规定精神学习教育总结报告', 'attrs': {'style': 'margin-top:21px;margin-bottom:21px;text-align:center;line-height:43px'}}, {'text': '为深入学习贯彻习近平新时代中国特色社会主义思想和党的二十大精神,坚决落实中央八项规定及其实施细则精神,按照上级党委的统一部署,XX市XX航空分公司党委自2025年X月起,在全公司范围内扎实开展了为期数月的中央八项规定精神学习教育。本次学习教育已于7月底圆满结束。现将公司学习教育的整体情况、存在 问题及未来规划总结汇报如下。', 'attrs': {'style': 'text-indent:43px;line-height:37px'}}, {'text': '一、学习教育开展情况', 'attrs': {'style': 'text-indent:43px;line-height:37px'}}, {'text': '公司党委坚持以习近平新时代中国特色社会主义思想为指导,把贯彻中央八项规定精神作为坚定拥护“两个确立”、坚决做到“两个维护”的重 大政治任务和具体行动。自学习教育启动以来,公司党委坚持高标准、严要求,精心组织、周密部署,确保学习教育有力有序有效推进。截至7月底,公司及所属各单位累计开展各类 专题学习研讨35次,通过线上线下相结合的方式,覆盖全体党员、干部及关键岗位人员1200余人次,实现重点人群100%全覆盖。在学习教育过程中,公司紧密结合航空运输企业特点, 聚焦突出问题,建立健全长效机制15项,累计查摆梳理出各类问题88个,并已全部完成整改,整改完成率达到100%。通过构建“学、查、改、督”四位一体的闭环管理机制,公司作风建 设取得显著成效,2025年上半年“三公”经费及非生产性支出同比下降18.5%,旅客及内部员工信访投诉数量同比减少22.3%。在实践中,公司探索并形成了“三个坚持”的宝贵经验:一是坚持高位推动与全员覆盖相结合,确保学习教育不留死角;二是坚持制度约束与文化涵养相结合,筑牢拒腐防变的思想和制度双重防线;三是坚持作风建设与安全发展相结合,以优良 作风为航空安全和高质量发展保驾护航。', 'attrs': {'style': 'text-indent:43px;line-height:37px'}}, {'text': '(一)学习研讨方面', 'attrs': {'style': 'text-indent:43px;line-height:37px'}}, {'text': '公司党委始终将理论武装置于首位,坚持在学深悟透中筑牢思想根基,推动学习教育从“有形覆盖”向“有效覆盖”深化。', 'attrs': {'style': 'text-indent:43px;line-height:37px'}}, {'text': '一是坚持领导带头,发挥“头雁效应”引领学。公司党委理论学习中心组率先垂范,先后组织了7次专题学习会,带头原 原本本、逐字逐句学习习近平总书记关于加强党的作风建设的重要论述、《锲而不舍落实中央八项规定精神,以优良党风引领社风民风》等重要文章以及新修订的《党政机关厉行 节约反对浪费条例》。公司党委书记、总经理等主要领导同志坚持先学一步、学深一层,并在学习会上结合分管领域和航空安全、市场营销、旅客服务等实际工作进行重点发言和 研讨交流,真正做到学思用贯通、知信行统一。这种自上而下的学习模式,为全体党员干部树立了标杆,有效激发了各级组织和广大党员的学习热情,形成了“领导领学、干部带学、 全员参学”的浓厚氛围。', 'attrs': {'style': 'text-indent:43px;line-height:37px'}}, {'text': '二是创新方式载体,推动“分层分类”精准学。针对公司组织层级多、专业 岗位杂、工作地点分散的特点,公司采取了差异化、精准化的学习策略。对于领导干部,重点深化对作风建设长期性、复杂性的认识,增强抓作风建设的政治自觉;对于新提拔的干部,专门开设“第一课”,强化纪律规矩教育,扣好廉洁从业的“第一粒扣子”;对于飞行、机务、运控等关键岗位人员,将学习内容与岗位职责、安全红线、职业操守紧密结合,强调规章制度的刚性执行;对于青年员工,通过举办青年理论学习小组读书班、知识竞赛、主题征文等形式,引导他们树立正确的价值观和事业观。同时,充分利用公司内网、“学习强国”APP、 企业微信群等线上平台,定期推送学习资料、典型案例和解读文章,将学习教育融入日常、抓在经常,有效解决了工学矛盾,确保了学习教育的全覆盖和高质量。', 'attrs': {'style': 'text-indent:43px;line-height:37px'}}, {'text': '三是强化警示教育,做到“以案促改”警钟长鸣学。公司党委深刻认识到,鲜活的案例是最好的清醒剂。学习教育期间,公司纪委系统梳理了党的十八大以来中央、民航系统及集团内部通报的违反中央八项规定精神的典型案例,特别是“四风”问题改头换面、隐蔽隐形的新表现新问题,汇编成《作风建设警示教育案例手册》,组织全体党员干部进行专题学习和对照反思。通过观看警示教育片、旁听职务犯罪庭审、参观XX市廉政教育基地等方式,用“身边事”教育“身边人”,引导广大 党员干部从思想深处受警醒、明底线、知敬畏。这种沉浸式、体验式的警示教育,使铁的纪律真正转化为党员干部的日常习惯和自觉遵循,达到了“查处一案、警示一片、治理一域”的良好效果,使“心有所畏、言有所戒、行有所止”成为普遍共识。', 'attrs': {'style': 'text-indent:43px;line-height:37px'}}, {'text': '(二)查摆问题方面', 'attrs': {'style': 'text-indent:43px;line-height:37px'}}, {'text': '公司坚持问题导向,以“刀刃向内”的勇气和“自我革命”的精神,深入查摆作风建设方面存在的顽瘴痼疾,确保问题找得准、挖得深、改得实。', 'attrs': {'style': 'text-indent:43px;line-height:37px'}}, {'text': '一是多维度对标对表,实现“立体式”检视。公司组织全体党员干部,特 别是领导班子成员,对照习近平总书记重要指示批示精神,对照党章党规党纪,对照中央八项规定及其实施细则精神,对照员工群众的期盼,深入开展“三对照三检视”活动。要求每位 党员干部把自己摆进去、把职责摆进去、把工作摆进去,认真撰写个人问题查摆清单。同时,公司党委主动对标行业内作风建设先进单位,查找在管理理念、制度执行、工作效能等 方面的差距,确保问题查摆既有思想层面的深度,也有工作实践中的精度。通过这一过程,共计梳理出涉及文山会海、检查考核过多过滥、服务基层意识不强等6大类共性问题。', 'attrs': {'style': 'text-indent:43px;line-height:37px'}}, {'text': '二是多渠道开门纳谏,实现“全景式”扫描。为确保问题找得全、找得实,公司畅通和拓宽了问题反映渠 道。一方面,通过召开不同层级、不同岗位的员工代表座谈会、发放无记名调查问卷、设立线上线下“作风建设意见箱”等方式,“面对面”“键对键”征求广大员工的意见建议。另一方面,结合“四下基层”制度,公司领导班子成员深入飞行分部、客舱部、维修基地、市场部等一线单位,与基层员工同学习、同劳动,在真诚交流中倾听心声、发现问题。学习教育期间,累计收集到来自基层一线的各类意见和建议210余条,为精准“画像”、靶向施治提供了第一手资料。', 'attrs': {'style': 'text-indent:43px;line-height:37px'}}, {'text': '三是高质量开展批评,实现“政治性”体检。在充分查摆和听取意见的基础上,公司党委和各基层党组织相继召开了高质量的专题民主生活会和组织生活会。会前,班子成员之间、 班子成员与分管部门负责人、党员与党员之间深入开展谈心谈话,把问题谈透、把思想谈通。会上,本着对党、对事业、对同志、对自已高度负责的精神,严肃认真开展批评和自我-批评,自我批评见人见事见思想,相互批评开诚布公、直截了当,真正起到了“红脸出汗、排毒治病”的效果。会后,针对查摆出的问题和相互批评的意见,建立动态更新的问题清单、 责任清单和整改清单“三张清单”,明确整改时限和责任人,实行挂图作战、销号管理,确保查摆出的每一个问题都得到有效解决。', 'attrs': {'style': 'text-indent:43px;line-height:37px'}}, {'text': '(三)集中整治方面', 'attrs': {'style': 'text-indent:43px;line-height:37px'}}, {'text': '公司坚持标本兼治、综合施策,针对查摆出的突出问题,开展专项整治行动,以“钉钉子”精神狠抓落实,推动作风建设取得实实在在的成效。', 'attrs': {'style': 'text-indent:43px;line-height:37px'}}, {'text': '一是紧盯顽瘴痼疾,开展“四风”问题专项整治。针对“四风”问题隐形变异、潜入地下的新动向,公司开展了为期三个月的专项整治行动。重点整治违规吃喝问题,严禁在内部食堂、培训中心 等场所搞“一桌餐”;重点整治违规收送礼品礼金问题,对利用电子红包、快递物流等方式收送“节礼”的行为进行重点排查;重点整治特权思想和特权现象,对公务舱改签、免票使用、贵宾休息室管理等环节的制度漏洞进行全面梳理和完善,严防利用公司资源谋取私利。通过专项整治,不仅刹住了一批歪风邪气,更重要的是完善了相关制度,扎紧了制度的笼子。', 'attrs': {'style': 'text-indent:43px;line-height:37px'}}, {'text': '二是聚焦主责主业,开展“作风建设与安全服务融合”专项整治。公司深刻认识到,航空企业的作风建 设必须与安全生产和旅客服务深度融合。为此,开展了“反形式主义、反官僚主义,促安全责任落实、促服务品质提升”的“两反两促”专项行动。在安全领域,重点整治安全检查走过 场、安全培训“雨过地皮湿”、隐患排查不深入不彻底等问题,通过实施领导干部带班检查、飞行和维修领域的“交叉检查、双人复核”等机制,有效提升了安全管理的精细化水平,今 年上半年公司安全责任原因万时率同比下降了0.05。在服务领域,重点整治对旅客诉求推诿扯皮、航班延误处置不力、服务流程僵化等“不担当、不作为”问题,通过优化旅客投诉处理流程,压缩处理时限,建立“首问负责制”,旅客满意度调查分数提升了3.2个百分点。', 'attrs': {'style': 'text-indent:43px;line-height:37px'}}, {'text': '三是力行勤 俭节约,开展“降本增效、过紧日子”专项整治。公司牢固树立“过紧日子”的思想,将厉行节约、反对浪费贯穿于经营管理全过程。严格执行“零基预算”管理,从严控制非生产性支出 。全面梳理办公用房、公务用车、会议培训、采购招标等环节,出台了《关于进一步加强成本管控的八条措施》,对差旅标准、招待费用、办公用品采购等进行刚性约束。通过技术创新和流程优化,在航材采购、燃油节约等方面深挖潜力。例如,通过优化航线、推广单发滑行等节油措施,上半年累计节约燃油成本近千万元。这种将作风建设内嵌于经营管理的 做法,不仅净化了风气,也实实在在地创造了经济效益,实现了作风建设与企业发展的双促进、双提升。', 'attrs': {'style': 'text-indent:43px;line-height:37px'}}, {'text': '(四)开门教育方面', 'attrs': {'style': 'text-indent:43px;line-height:37px'}}, {'text': '公司坚持走好新时代党的群众路线,把“开门”贯穿于学习教育全过程,主动 接受员工群众监督,问计于民、问需于民,使学习教育的过程成为密切联系群众、凝聚发展合力的过程。', 'attrs': {'style': 'text-indent:43px;line-height:37px'}}, {'text': '一是领导干部“沉下去”,深入基层听民意。严格落实领导干部“四下基层”制度,公司领导班子成员每人确定1-2个基层联系点,定期深入一线开展工作。不发通知、不打招呼,直奔基层、直插现场,与飞行员、乘务员、机务维修师、地服人员等一线员工同坐一条板凳,倾听他们的烦心事、揪心事。在了解到一线员工对于倒班休息、通勤交通、子女入学等方面的困难后,公司党委立即召开专题会议研究,推动解决了员工公寓空调增配、优化夜间通勤班车线路等一批“急难愁盼”问题,让员工切实感受到了组织的关怀和温暖。', 'attrs': {'style': 'text-indent:43px;line-height:37px'}}, {'text': '二是开展“换位体验”活动,设身处地察实情。为打破部门壁垒,根治机关作风中存在的官僚主义,公司创新性地 组织机关职能部门的干部定期到一线岗位进行“换位体验”。财务部的干部到售票柜台体验报销流程的繁琐,人力资源部的干部到客舱部体验乘务员的辛劳,通过亲身体验,使管理者 能够真正站在一线员工和旅客的角度思考问题、改进工作。这一活动有效促进了机关服务意识的转变,推动了多个跨部门业务流程的优化再造,办事效率平均提升了15%以上。', 'attrs': {'style': 'text-indent:43px;line-height:37px'}}, {'text': '三是搭建监督互动平台,广开言路纳良策。公司充分发挥职工代表大会、工会等组织的作用,定期召开恳谈会,通报公司作风建设情况,听取意见建议。同时,在公司内网开设“作风建设直通车”专栏,员工可以匿名反映问题、提出建议。公司纪委和相关职能部门对专栏反映的问题建立台账,限时办理、及时反馈。学习教育期间,通过该平台采纳并实施了关于“优化机组过夜酒店标准”“简化航材申领流程”等合理化建议30余项,不仅解决了实际问题,更营造了民主开放、群策群力的良好政治生态。', 'attrs': {'style': 'text-indent:43px;line-height:37px'}}, {'text': '(五)组织领导方面', 'attrs': {'style': 'text-indent:43px;line-height:37px'}}, {'text': '公司党委从始至终高度重视本次学习教育,将其作为一项重大政治任务来抓,加强组织领导,压实各级责任,强化督促检查,为学习教育的顺利开展提 供了坚强的组织保障。', 'attrs': {'style': 'text-indent:43px;line-height:37px'}}, {'text': '一是健全领导机制,构建“责任共同体”。公司第一时间成立了由党委书记任组长、纪委书记任副组长的学习教育领导小组,下设办公室,负责统筹协调和日常工作。制定了详尽的实施方案,明确了学习教育的指导思想、目标任务、方法步骤和工作要求。形 成了“党委统一领导、领导小组牵头抓总、纪委监督推动、职能部门协同配合、基层党组织具体落实”的领导体制和工作格局,确保各项任务有人抓、有人管、能落实。', 'attrs': {'style': 'text-indent:43px;line-height:37px'}}, {'text': '二是压实主体责任,拧紧“一级抓一级”的责任链条。公司党委坚决扛起主体责任,党委书记认真履行“第一责任 人”职责,对重要工作亲自部署、重大问题亲自过问、重要环节亲自协调。领导班子成员严格落实“一岗双责”,既抓好分管业务工作,又抓好分管领域和部门的作风建设。各基层党组织书记切实担负起直接责任,把学习教育抓在手上、扛在肩上。通过层层签订责任书、逐级传导压力,构建起横向到边、纵向到底的责任体系,有效防止了“上热中温下冷”现象。', 'attrs': {'style': 'text-indent:43px;line-height:37px'}}, {'text': '三是强化督导检查,用好“考核指挥棒”。领导小组办公室联合公司纪委、党群工作部等部门组成督导 组,采取“四不两直”的方式,对各单位学习教育开展情况进行全过程、滚动式督导检查。将学习教育的成效与各单位的年度绩效考核、评优评先以及干部选拔任用直接挂钩,对思想 上不重视、工作上不得力的单位和个人,及时约谈提醒、严肃追责问责。这种动真碰硬的督导考核,确保了学习教育不虚、不空、不走过场,推动了中央八项规定精神在公司落地生 根。', 'attrs': {'style': 'text-indent:43px;line-height:37px'}}, {'text': '二、存在问题', 'attrs': {'style': 'text-indent:43px;line-height:37px'}}, {'text': '在总结成绩和经验的同时,公司党委也清醒地认识到,对照中央、上级党委的高标准、严要求,对照员工群众的新期待,本次学习教育和公司常态化的作风建设仍存在一些不容忽视的问题和短板,主要体现在以下几个方面:', 'attrs': {'style': 'text-indent:43px;line-height:37px'}}, {'text': '(一)学习教育的深度和广度有待进一步拓展。部分基层 单位和少数党员干部对学习教育的极端重要性认识仍不够深刻,存在一定的“闯关”思想和“应付”心态。学习研讨有时满足于“读过了、学过了”,但在“学懂、弄通、做实”上下的功夫还不够,理论联系实际、指导实践的能力有待加强。学习教育的覆盖面虽然广,但在一些业务繁忙的一线班组,学习的系统性和持续性保障不足,存在“前紧后松”的现象。', 'attrs': {'style': 'text-indent:43px;line-height:37px'}}, {'text': '(二)纠治“四风”问题的韧劲和精度有待进一步提升。随着作风建设的不断深入,“四风”问题变得更加隐蔽、更具迷惑性。例如,通过电子支付方式收送红包、礼金,以“土特产”名义进行利益输送,将公务接待转嫁给关联单位等现象偶有发生,发现和查处的难度增大。个别领域的形式主义、官僚主义问题依然存在,如一些会议文件仍然偏多偏长,一些检查考核存在重形式、轻实效的倾向,需要下更大力气进行根治。', 'attrs': {'style': 'text-indent:43px;line-height:37px'}}, {'text': '(三)制度体系的系统性和协同性有待进一步完善。尽管在学习教育中建立和完善了一批制度,但对照新形势新要求,部分现有制度仍然存在滞后性和“空白 点”。例如,在供应商管理、航线合作、大宗采购等高风险领域,制度规定还不够细化,操作性有待增强。各制度之间的衔接不够紧密,存在“制度打架”或协同效应不强的问题,未能完全形成靠制度管权、管事、管人的闭环体系。', 'attrs': {'style': 'text-indent:43px;line-height:37px'}}, {'text': '(四)常态化长效化机制的执行力有待进一步强化。 作风建设非一日之功,必须常抓不懈。当前存在的一种倾向是,在集中学习教育期间抓得紧、抓得严,但活动过后可能出现思想松懈、力度减弱的风险。监督检查的常态化、制度化 还有待加强,日常监督的“探头”作用发挥还不够充分,如何将集中整治的成果巩固好、发展好,防止问题反弹回潮,是需要持续用力解决的课题。', 'attrs': {'style': 'text-indent:43px;line-height:37px'}}, {'text': '三、下一步打算', 'attrs': {'style': 'text-indent:43px;line-height:37px'}}, {'text': '针对上述存在的问题,公司党委将坚 持目标不变、标准不降、力度不减,以此次学习教育总结为新的起点,持续发力、久久为功,推动作风建设向纵深发展。', 'attrs': {'style': 'text-indent:43px;line-height:37px'}}, {'text': '(一)在深化理论武装上持续用力,推动思想认识再提升。坚持把学习贯彻习近平新时代中国特色社会主义思想作为首要政治任务,建立理论学习常态化机制,将 中央八项规定精神作为党委理论学习中心组、各级党组织“三会一课”、干部培训的必修课、常修课。创新学习方式,多运用案例式、研讨式、情景式教学,增强学习的吸引力和实效性。持续强化警示教育,定期通报内外部违纪违法典型案例,教育引导全体党员干部筑牢信仰之基、补足精神之钙、把稳思想之舵。', 'attrs': {'style': 'text-indent:43px;line-height:37px'}}, {'text': '(二)在完善制度体系上持续用力,推动源头治理再加强。启动公司内部规章制度的全面梳理工作,按照“废、改、立”的原则,对与中央精神和现实需要不符的制度坚决予以废止,对不完善的制度抓紧修订,对制度空白点及时予以填补。重点围绕公务接待、差旅管理、物资采购、营销费用、员工福利等关键环节,制定更加具体、 更具操作性的实施细则。加强制度执行的协同性,打通不同制度之间的堵点,构建系统完备、科学规范、运行有效的作风建设制度体系,真正实现用制度管人管事。', 'attrs': {'style': 'text-indent:43px;line-height:37px'}}, {'text': '(三)在强化监督执纪上持续用力,推动纪律约束再绷紧。坚持严的主基调不动摇,对违反中央八项规定精神的问题, 无论涉及到谁,都坚持“零容忍”的态度,发现一起、查处一起、通报一起。加强纪检、审计、财务、人事等部门的协作联动,运用大数据等信息化手段,对费用报销、采购招标、项目建设等领域进行精准监督,提升发现问题的能力。强化日常监督,将监督的触角延伸到最基层,综合运用谈话提醒、批评教育、诫勉函询等方式,抓早抓小、防微杜渐,让咬耳扯袖、 红脸出汗成为常态。', 'attrs': {'style': 'text-indent:43px;line-height:37px'}}, {'text': '(四)在深化标本兼治上持续用力,推动作风建设再融合。坚持将作风建设与公司改革发展稳定大局同谋划、同部署、同推进。将落实中央八项规定精神的要求深度融入到航空安全管理、航班运行保障、旅客服务提升、市场经营开拓等各项业务工作中,以过 硬的作风保障绝对安全,以务实的作风提升服务品质,以清廉的作风促进效益增长。大力弘扬新时代民航精神和公司优秀企业文化,选树宣传一批作风优良、业绩突出的先进典型,营造风清气正、干事创业的良好政治生态,以作风建设的新成效为公司在XX区域的高质量发展提供坚强保障,为集团公司建设世界一流航空企业作出新的更大贡献。', 'attrs': {'style': 'text-indent: 43px; line-height: 37px;'}}] # 你的数据
    out_path="20250907/XX市XX航空分公司深入贯彻中央八项规定精神学习教育总结报告.docx"
    formattedTextObjects = formated_text_object(text_objects)
    # print("formattedTextObjects",formattedTextObjects)
    out_path = "20250907/XX市XX航空分公司深入贯彻中央八项规定精神学习教育总结报告.docx"
    save_to_word(formattedTextObjects,out_path)
